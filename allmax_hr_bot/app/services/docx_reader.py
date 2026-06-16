from __future__ import annotations

from functools import lru_cache
from pathlib import Path
import re
from docx import Document

from app.config import settings
from app.utils.texts import POSITIONS

POSITION_FILES = {
    "Merchendayzer-konsultant": ["Savdo reglamenti"],
    "WMS xodimi": ["WMS.docx", "WMS mikro", "Omborxona Xavfsizlik", "MAHSULOT NOMLASH"],
    "Dazmollash bo‘limi": ["WMS mikro", "Omborxona Xavfsizlik"],
    "Storis Maker": ["STORIS", "Ijtimoiy tarmoq", "MAHSULOT NOMLASH"],
    "Mobilograf": ["mobilograf", "Ijtimoiy tarmoq", "MAHSULOT NOMLASH"],
    "Dizayner": ["grafik dizayner", "Ijtimoiy tarmoq", "MAHSULOT NOMLASH"],
    "Brand Face": ["BRAND FACE", "Ijtimoiy tarmoq"],
    "Kassa": ["Kassa"],
    "Community Manager": ["Community manager", "Ijtimoiy tarmoq", "MAHSULOT NOMLASH"],
    "Call Centre / Call operator": ["Call operator"],
    "Copywriter": ["Copywriter", "Ijtimoiy tarmoq", "MAHSULOT NOMLASH"],
}

COPYWRITER_FALLBACK = """
Copywriter uchun universal reglament:
- mahsulot va brend uchun matn yozish;
- reklama postlari, story matnlari va scriptlar tayyorlash;
- mijozni jalb qiluvchi qisqa, aniq va sotuvga yo‘naltirilgan matn yozish;
- grammatik to‘g‘rilik, brend ohangi va mahsulot nomlash standartiga rioya qilish;
- deadline, tartib-intizom va yakuniy matn sifatini shaxsan tekshirish;
- AI yoki boshqa vositalardan foydalansa ham final sifat uchun javobgarlik.
"""


def _read_docx(path: Path) -> str:
    try:
        doc = Document(path)
        parts: list[str] = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        return f"[DOCX o‘qishda xato: {path.name}: {exc}]"


@lru_cache(maxsize=1)
def available_files() -> list[Path]:
    return [p for p in settings.regulations_dir.rglob("*.docx") if p.is_file()]


def _matches(path: Path, pattern: str) -> bool:
    name = path.name.lower()
    pattern = pattern.lower().replace(".docx", "")
    return pattern in name or re.sub(r"\s+", " ", pattern) in re.sub(r"\s+", " ", name)


@lru_cache(maxsize=64)
def get_regulation_text(position: str) -> str:
    patterns = POSITION_FILES.get(position, [])
    files: list[Path] = []
    for pat in patterns:
        for p in available_files():
            if _matches(p, pat) and p not in files:
                files.append(p)
    sections: list[str] = []
    for p in files:
        sections.append(f"\n=== {p.name} ===\n{_read_docx(p)}")
    if position == "Copywriter" and not sections:
        sections.append(COPYWRITER_FALLBACK)
    if not sections:
        sections.append(f"{position} uchun reglament topilmadi. Umumiy ALLMAX qadriyatlari: mas’uliyat, halollik, intizom, mijozga hurmat, o‘rganishga tayyorlik.")
    text = "\n".join(sections)
    return text[:60000]


def brief_for_position(position: str) -> str:
    text = get_regulation_text(position)
    lines = [ln.strip("•- ").strip() for ln in text.splitlines() if len(ln.strip()) > 15]
    compact = []
    seen = set()
    for ln in lines:
        key = ln.lower()[:80]
        if key not in seen:
            compact.append(ln)
            seen.add(key)
        if len(compact) >= 8:
            break
    main = "\n".join(f"• {x}" for x in compact[:6]) or "• Mas’uliyat\n• Tartib-intizom\n• Mijozga hurmat"
    return (
        f"<b>{position}</b>\n\n"
        f"<b>Asosiy vazifalar va talablar:</b>\n{main}\n\n"
        "<b>Stajirovka:</b> 7 kunlik o‘quv jarayoni, 20 dars, har darsdan keyin 10 talik test va yakuniy 30 talik test mavjud."
    )


def all_positions_text() -> str:
    return "\n".join(f"{i}. {p}" for i, p in enumerate(POSITIONS, 1))
