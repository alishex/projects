from __future__ import annotations

import hashlib
import json
import re
import shutil
from datetime import datetime
from pathlib import Path
from uuid import uuid4

from docx import Document
from pypdf import PdfReader

from app.config import settings
from app.database import db, loads
from app.utils.texts import POSITIONS

DEFAULT_VACANCIES = {
    "Merchendayzer-konsultant": ["Savdo reglamenti"],
    "WMS xodimi": ["WMS.docx", "WMS mikro", "Omborxona Xavfsizlik", "MAHSULOT NOMLASH"],
    "Dazmollash bo‘limi": ["WMS mikro", "Omborxona Xavfsizlik"],
    "Storis Maker": ["STORIS", "Ijtimoiy tarmoq", "MAHSULOT NOMLASH"],
    "Mobilograf": ["mobilograf", "Ijtimoiy tarmoq", "MAHSULOT NOMLASH"],
    "Dizayner": ["grafik dizayner", "Ijtimoiy tarmoq", "MAHSULOT NOMLASH"],
    "Brand Face": ["BRAND FACE", "Ijtimoiy tarmoq"],
    "Kassa": ["Kassa"],
    "Community Manager": ["Community manager", "Ijtimoiy tarmoq", "MAHSULOT NOMLASH"],
    "Call Centre / Call operator": ["Call operator"],
    "Copywriter": ["Ijtimoiy tarmoq", "MAHSULOT NOMLASH"],
}

COPYWRITER_FALLBACK = """Copywriter uchun universal reglament:
- mahsulot va brend uchun matn yozish;
- reklama postlari, story matnlari va scriptlar tayyorlash;
- mijozni jalb qiluvchi qisqa, aniq va sotuvga yo‘naltirilgan matn yozish;
- grammatik to‘g‘rilik, brend ohangi va mahsulot nomlash standartiga rioya qilish;
- deadline, tartib-intizom va yakuniy matn sifatini shaxsan tekshirish;
- AI yoki boshqa vositalardan foydalansa ham final sifat uchun javobgarlik.
"""

REGULATION_TYPES = [
    "Asosiy lavozim reglamenti",
    "Xavfsizlik reglamenti",
    "Umumiy kompaniya qadriyatlari",
    "Ish grafigi reglamenti",
    "Mahsulot standarti",
    "Kontent standarti",
    "Qo‘shimcha qoida",
]


def safe_filename(filename: str) -> str:
    path = Path(filename or "reglament.txt")
    stem = re.sub(r"[^A-Za-z0-9А-Яа-яЁёЎўҚқҒғҲҳ _.-]+", "_", path.stem).strip(" ._") or "reglament"
    stem = re.sub(r"\s+", "_", stem)[:80]
    ext = path.suffix.lower() if path.suffix.lower() in {".docx", ".pdf", ".txt"} else ".txt"
    return f"{stem}_{uuid4().hex[:10]}{ext}"


def extract_text(path: Path) -> tuple[str, str | None]:
    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            doc = Document(path)
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts).strip()
        elif suffix == ".pdf":
            reader = PdfReader(str(path))
            text = "\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()
            if len(text) < 80:
                return text, "Ushbu PDF ichidagi matn avtomatik to‘liq o‘qilmadi. Iltimos, DOCX yoki TXT formatda yuboring."
        elif suffix == ".txt":
            text = path.read_text(encoding="utf-8", errors="replace").strip()
        else:
            return "", "Faqat .docx, .pdf yoki .txt fayl yuboring."
    except Exception as exc:
        return "", f"Fayl matnini o‘qishda xato: {exc}"
    if len(text) < 40:
        return text, "Reglament matni yetarli o‘qilmadi. DOCX yoki TXT formatda to‘liq fayl yuboring."
    return text, None


def save_uploaded_document(source_path: Path, original_name: str) -> tuple[Path, str, str | None]:
    upload_dir = settings.regulations_dir / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)
    name = safe_filename(original_name)
    target = upload_dir / name
    shutil.copy2(source_path, target)
    text, warning = extract_text(target)
    return target, text, warning


def summarize_text(text: str, max_lines: int = 7) -> str:
    useful = []
    seen = set()
    for raw in text.splitlines():
        line = raw.strip(" •-\t")
        if len(line) < 12:
            continue
        key = line.lower()[:80]
        if key not in seen:
            useful.append(line)
            seen.add(key)
        if len(useful) >= max_lines:
            break
    return "\n".join(f"• {line}" for line in useful) or "• Matndan aniq bandlar ajratilmadi."


def compare_regulations(old_text: str, new_text: str) -> str:
    old_lines = {x.strip() for x in old_text.splitlines() if len(x.strip()) > 20}
    new_lines = {x.strip() for x in new_text.splitlines() if len(x.strip()) > 20}
    added = list(new_lines - old_lines)[:5]
    removed = list(old_lines - new_lines)[:5]
    lines = []
    if added:
        lines.append("Yangi yoki o‘zgargan bandlar:")
        lines += [f"+ {x}" for x in added]
    if removed:
        lines.append("Avvalgi versiyada bo‘lgan, yangi matnda topilmagan bandlar:")
        lines += [f"- {x}" for x in removed]
    return "\n".join(lines) or "Matnlar orasida asosiy farq avtomatik aniqlanmadi."


def _matches(path: Path, pattern: str) -> bool:
    return pattern.lower().replace(".docx", "") in path.name.lower()


def bootstrap_dynamic_catalog() -> None:
    """Seed old shipped catalog into DB once; afterwards admin manages everything in Telegram."""
    for pos in POSITIONS:
        if not db.get_vacancy_by_name(pos):
            db.create_vacancy({
                "name_uz": pos,
                "name_ru": pos,
                "description_uz": f"{pos} lavozimi bo‘yicha ish o‘rni.",
                "responsibilities": "Reglament asosidagi vazifalarni o‘z vaqtida va sifatli bajarish.",
                "requirements": "Mas’uliyat, intizom, halollik va o‘rganishga tayyorlik.",
                "work_schedule": "Reglament yoki HR bilan kelishuv asosida",
                "status": "active",
            }, None)
    existing = db.list_regulations()
    if existing:
        return
    file_to_reg_id: dict[str, int] = {}
    shipped_files = [p for p in settings.regulations_dir.glob("*.docx") if p.is_file()]
    for file_path in shipped_files:
        text, warning = extract_text(file_path)
        if warning and not text:
            continue
        title = file_path.stem
        rtype = "Xavfsizlik reglamenti" if "xavfsizlik" in title.lower() else "Asosiy lavozim reglamenti"
        rid = db.create_regulation(title, rtype, None)
        vid = db.add_regulation_version(rid, file_path.name, str(file_path.relative_to(settings.regulations_dir.parent)), ".docx", text, summarize_text(text), "Boshlang‘ich import", None, active=True)
        file_to_reg_id[file_path.name] = rid
    for position, patterns in DEFAULT_VACANCIES.items():
        vacancy = db.get_vacancy_by_name(position)
        if not vacancy:
            continue
        for path in shipped_files:
            if any(_matches(path, pat) for pat in patterns):
                title = path.stem
                matches = [r for r in db.list_regulations() if r["title"] == title]
                if matches:
                    db.link_vacancy_regulation(int(vacancy["id"]), int(matches[0]["id"]), None)
        if position == "Copywriter" and not db.vacancy_regulation_links(int(vacancy["id"])):
            rid = db.create_regulation("Copywriter universal reglamenti", "Asosiy lavozim reglamenti", None)
            vid = db.add_regulation_version(rid, "Copywriter_fallback.txt", "generated", ".txt", COPYWRITER_FALLBACK, summarize_text(COPYWRITER_FALLBACK), "Universal fallback", None, active=True)
            db.link_vacancy_regulation(int(vacancy["id"]), rid, None, vid)


def snapshot_items(snapshot: str | None) -> list[dict]:
    return loads(snapshot, []) if snapshot else []


def regulation_text_for_vacancy(vacancy_id: int | None, purpose: str = "interview", snapshot: str | None = None) -> str:
    if not vacancy_id:
        return "Reglament biriktirilmagan. Admin qo‘shimcha reglament yuklashi kerak."
    sections: list[str] = []
    if snapshot:
        items = snapshot_items(snapshot)
        for item in items:
            use_key = {"interview": "use_for_interview", "lessons": "use_for_lessons", "tests": "use_for_tests"}.get(purpose)
            if use_key and not int(item.get(use_key, 1)):
                continue
            version = db.get_regulation_version(int(item.get("version_id") or 0))
            if version:
                sections.append(f"=== {version.get('title')} — v{version.get('version_number')} ===\n{version.get('extracted_text') or ''}")
    else:
        for link in db.vacancy_regulation_links(vacancy_id, purpose):
            sections.append(f"=== {link.get('title')} — v{link.get('version_number')} ===\n{link.get('extracted_text') or ''}")
    if not sections:
        vacancy = db.get_vacancy(vacancy_id) or {}
        return f"{vacancy.get('name_uz') or 'Vakansiya'} uchun faol reglament ma’lumoti yetarli emas. Qo‘shimcha reglament yuklang."
    return "\n\n".join(sections)[:90000]


def brief_for_vacancy(vacancy: dict, lang: str = "uz") -> str:
    name = vacancy.get("name_ru") if lang == "ru" and vacancy.get("name_ru") else vacancy.get("name_uz")
    desc = vacancy.get("description_ru") if lang == "ru" and vacancy.get("description_ru") else vacancy.get("description_uz")
    return (
        f"<b>{name}</b>\n\n"
        f"<b>Qisqa ma’lumot:</b>\n{desc or '-'}\n\n"
        f"<b>Asosiy vazifalar:</b>\n{vacancy.get('responsibilities') or '-'}\n\n"
        f"<b>Asosiy talablar:</b>\n{vacancy.get('requirements') or '-'}\n\n"
        f"<b>Ish grafigi:</b> {vacancy.get('work_schedule') or '-'}\n"
        f"<b>Stajirovka:</b> {'Mavjud' if vacancy.get('internship_enabled') else 'Mavjud emas'}"
        f" ({vacancy.get('internship_days') or 7} kun, {vacancy.get('lesson_count') or 20} dars, yakuniy {vacancy.get('final_test_count') or 30} talik test)"
    )


def catalog_warnings() -> list[str]:
    warnings: list[str] = []
    for vacancy in db.list_vacancies(status="active"):
        vid = int(vacancy["id"])
        name = vacancy.get("name_uz") or "Vakansiya"
        links = db.vacancy_regulation_links(vid)
        if not links:
            warnings.append(f"{name} uchun faol reglament hali biriktirilmagan.")
            continue
        has_active_lessons = bool(db.list_training_materials(vid, "active", "lesson", 1))
        if vacancy.get("internship_enabled") and not has_active_lessons:
            warnings.append(f"{name} uchun dars materiallari hali active qilinmagan; boshlanganda AI/fallback yaratiladi.")
    return warnings[:12]
